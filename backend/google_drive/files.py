"""Utilities for fetching and transforming Google Drive files for RAG."""

from __future__ import annotations

import io
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Dict, List, Optional, Sequence, Tuple

import docx
import httpx
import pandas as pd
import tiktoken
from langchain_core.documents import Document


FILES_ENDPOINT = "https://www.googleapis.com/drive/v3/files"

_DEFAULT_FIELDS = (
    "nextPageToken, files(id, name, mimeType, modifiedTime, webViewLink,"
    " size, capabilities/canDownload)"
)

_EXPORT_MIME_MAP: Dict[str, Optional[str]] = {
    "application/vnd.google-apps.document": "text/plain",
    "application/vnd.google-apps.spreadsheet": "text/csv",
    "application/vnd.google-apps.presentation": "text/plain",
}

_TEXT_MIME_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "application/csv",
    "application/json",
}

_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

_XLSX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}

_PDF_MIME_TYPES = {"application/pdf"}

_HWP_MIME_TYPES = {"application/x-hwp", "application/haansoft-hwp"}

_ENCODINGS = ("utf-8", "cp949", "euc-kr", "latin-1")

_ENC = tiktoken.get_encoding("cl100k_base")

_DEFAULT_CHUNK_SIZE = 800
_DEFAULT_CHUNK_OVERLAP_RATIO = 0.1


class GoogleDriveAPIError(RuntimeError):
    """Raised when the Google Drive API responds with an unexpected error."""


class UnsupportedGoogleDriveFile(RuntimeError):
    """Raised when a Google Drive file type cannot be processed."""


@dataclass(slots=True)
class GoogleDriveFile:
    """Container for the textual representation of a Google Drive file."""

    file_id: str
    name: str
    mime_type: str
    modified_time: str
    web_view_link: Optional[str]
    text: str
    format: str


def _format_datetime_for_query(value: datetime) -> str:
    if value.tzinfo is None:
        value = value.replace(tzinfo=timezone.utc)
    else:
        value = value.astimezone(timezone.utc)
    return value.isoformat().replace("+00:00", "Z")


async def _list_files(
    client: httpx.AsyncClient,
    *,
    access_token: str,
    modified_after: Optional[datetime] = None,
) -> List[Dict[str, str]]:
    headers = {"Authorization": f"Bearer {access_token}"}
    params = {
        "pageSize": 200,
        "fields": _DEFAULT_FIELDS,
        "supportsAllDrives": "false",
        "includeItemsFromAllDrives": "false",
        "orderBy": "modifiedTime desc",
        "q": " and ".join(
            filter(
                None,
                [
                    "trashed=false",
                    "mimeType != 'application/vnd.google-apps.folder'",
                    "'me' in owners",
                    (
                        f"modifiedTime > '{_format_datetime_for_query(modified_after)}'"
                        if modified_after
                        else None
                    ),
                ],
            )
        ),
    }

    files: List[Dict[str, str]] = []
    page_token: Optional[str] = None

    while True:
        if page_token:
            params["pageToken"] = page_token
        else:
            params.pop("pageToken", None)

        response = await client.get(FILES_ENDPOINT, headers=headers, params=params)
        if response.status_code == 401:
            raise GoogleDriveAPIError("Google Drive 접근 권한이 만료되었습니다.")
        if response.status_code != 200:
            raise GoogleDriveAPIError(response.text)

        payload = response.json()
        files.extend(payload.get("files", []))
        page_token = payload.get("nextPageToken")
        if not page_token:
            break

    return files


def _decode_text(content: bytes) -> str:
    for encoding in _ENCODINGS:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _convert_docx(content: bytes) -> str:
    document = docx.Document(io.BytesIO(content))
    lines = [paragraph.text for paragraph in document.paragraphs]
    return "\n".join(filter(None, lines)).strip()


def _convert_excel(content: bytes) -> str:
    data = io.BytesIO(content)
    frames = pd.read_excel(data, sheet_name=None, dtype=str, na_filter=False)
    parts: List[str] = []
    for sheet_name, frame in frames.items():
        parts.append(f"# Sheet: {sheet_name}")
        parts.append(frame.to_csv(index=False))
    return "\n\n".join(parts).strip()


def _chunk_text(
    text: str,
    *,
    chunk_size: int = _DEFAULT_CHUNK_SIZE,
    overlap_ratio: float = _DEFAULT_CHUNK_OVERLAP_RATIO,
) -> List[str]:
    if not text:
        return []

    tokens = _ENC.encode(text)
    if not tokens:
        return []

    overlap = int(chunk_size * overlap_ratio)
    if overlap >= chunk_size:
        overlap = max(0, chunk_size - 1)

    chunks: List[str] = []
    start = 0
    total = len(tokens)

    while start < total:
        end = min(total, start + chunk_size)
        chunk_tokens = tokens[start:end]
        chunks.append(_ENC.decode(chunk_tokens))
        if end >= total:
            break
        start = end - overlap if overlap > 0 else end

    return chunks


def _build_records_from_file(
    file: GoogleDriveFile,
    *,
    chunk_size: int = _DEFAULT_CHUNK_SIZE,
    overlap_ratio: float = _DEFAULT_CHUNK_OVERLAP_RATIO,
) -> List[Dict[str, str]]:
    records: List[Dict[str, str]] = []
    chunks = _chunk_text(file.text, chunk_size=chunk_size, overlap_ratio=overlap_ratio)

    if not chunks:
        records.append(
            {
                "file_id": file.file_id,
                "title": file.name,
                "modified_time": file.modified_time,
                "url": file.web_view_link or "",
                "mime_type": file.mime_type,
                "text": "",
                "plain_text": "",
                "format": file.format,
                "chunk_index": 0,
            }
        )
        return records

    for chunk_index, chunk in enumerate(chunks):
        records.append(
            {
                "file_id": file.file_id,
                "title": file.name,
                "modified_time": file.modified_time,
                "url": file.web_view_link or "",
                "mime_type": file.mime_type,
                "text": chunk,
                "plain_text": chunk,
                "format": file.format,
                "chunk_index": chunk_index,
            }
        )

    return records


def build_records_from_files(
    files: Sequence[GoogleDriveFile],
    *,
    chunk_size: int = _DEFAULT_CHUNK_SIZE,
    overlap_ratio: float = _DEFAULT_CHUNK_OVERLAP_RATIO,
) -> List[Dict[str, str]]:
    records: List[Dict[str, str]] = []
    for file in files:
        records.extend(
            _build_records_from_file(
                file, chunk_size=chunk_size, overlap_ratio=overlap_ratio
            )
        )
    return records


def build_documents_from_records(
    records: Sequence[Dict[str, str]],
    workspace_metadata: Dict[str, str],
) -> List[Document]:
    documents: List[Document] = []
    for index, record in enumerate(records):
        plain_text = record.get("plain_text") or ""
        formatted_text = record.get("text") or plain_text
        if not plain_text.strip():
            continue

        metadata = dict(workspace_metadata)
        file_id = record.get("file_id") or ""
        chunk_index = record.get("chunk_index")
        if isinstance(chunk_index, int):
            chunk_number = chunk_index
        else:
            chunk_number = index
        metadata.update(
            {
                "page_id": file_id,
                "page_title": record.get("title"),
                "page_url": record.get("url"),
                "last_edited_time": record.get("modified_time"),
                "chunk_id": f"{file_id}:{chunk_number}",
                "chunk_index": chunk_number,
                "format": record.get("format", "text"),
                "formatted_text": formatted_text,
                "plain_text": plain_text,
                "provider": "googledrive",
                "file_id": file_id,
                "file_mime_type": record.get("mime_type"),
            }
        )
        documents.append(Document(page_content=plain_text, metadata=metadata))
    return documents


async def _download_file(
    client: httpx.AsyncClient,
    *,
    file: Dict[str, str],
    headers: Dict[str, str],
) -> Tuple[str, str]:
    file_id = file.get("id")
    mime_type = file.get("mimeType") or ""
    name = file.get("name") or ""

    if mime_type in _HWP_MIME_TYPES or name.lower().endswith(".hwp"):
        raise UnsupportedGoogleDriveFile("한글(.hwp) 파일은 지원하지 않습니다.")

    export_mime = _EXPORT_MIME_MAP.get(mime_type)
    if export_mime is not None:
        params = {"mimeType": export_mime}
        response = await client.get(
            f"{FILES_ENDPOINT}/{file_id}/export",
            headers=headers,
            params=params,
        )
        if response.status_code != 200:
            raise GoogleDriveAPIError(response.text)
        content = response.content
        mime_type = export_mime
    else:
        response = await client.get(
            f"{FILES_ENDPOINT}/{file_id}",
            headers=headers,
            params={"alt": "media"},
        )
        if response.status_code != 200:
            raise GoogleDriveAPIError(response.text)
        content = response.content

    if mime_type in _TEXT_MIME_TYPES:
        return _decode_text(content), "text"
    if mime_type in _DOCX_MIME_TYPES:
        return _convert_docx(content), "text"
    if mime_type in _XLSX_MIME_TYPES:
        return _convert_excel(content), "csv"
    if mime_type in _PDF_MIME_TYPES:
        raise UnsupportedGoogleDriveFile("PDF 파일은 현재 지원하지 않습니다.")

    raise UnsupportedGoogleDriveFile(
        f"지원하지 않는 Google Drive 파일 형식입니다: {mime_type}"
    )


async def fetch_authorized_text_files(
    access_token: str,
    *,
    modified_after: Optional[datetime] = None,
) -> Tuple[List[GoogleDriveFile], List[Dict[str, str]]]:
    async with httpx.AsyncClient(timeout=60) as client:
        raw_files = await _list_files(
            client, access_token=access_token, modified_after=modified_after
        )

        headers = {"Authorization": f"Bearer {access_token}"}
        converted: List[GoogleDriveFile] = []
        skipped: List[Dict[str, str]] = []

        for file in raw_files:
            capabilities = file.get("capabilities") or {}
            can_download = capabilities.get("canDownload", True)
            if not can_download and file.get("mimeType") not in _EXPORT_MIME_MAP:
                skipped.append(
                    {
                        "file_id": file.get("id"),
                        "name": file.get("name"),
                        "mime_type": file.get("mimeType"),
                        "reason": "다운로드 권한이 없습니다.",
                    }
                )
                continue

            try:
                text, fmt = await _download_file(
                    client, file=file, headers=headers
                )
            except UnsupportedGoogleDriveFile as exc:
                skipped.append(
                    {
                        "file_id": file.get("id"),
                        "name": file.get("name"),
                        "mime_type": file.get("mimeType"),
                        "reason": str(exc),
                    }
                )
                continue
            except GoogleDriveAPIError as exc:
                skipped.append(
                    {
                        "file_id": file.get("id"),
                        "name": file.get("name"),
                        "mime_type": file.get("mimeType"),
                        "reason": f"API 오류: {exc}",
                    }
                )
                continue

            converted.append(
                GoogleDriveFile(
                    file_id=file.get("id", ""),
                    name=file.get("name", ""),
                    mime_type=file.get("mimeType", ""),
                    modified_time=file.get("modifiedTime", ""),
                    web_view_link=file.get("webViewLink"),
                    text=text,
                    format=fmt,
                )
            )

    return converted, skipped
